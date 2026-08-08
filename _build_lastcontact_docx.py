from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn


ROOT = Path(r"C:\Nyasha_Planet\git\lastcontact")
TEMPLATE_DOCX = ROOT / "LastContact_autoWrap.docx"
TXT_PATH = ROOT / "LastContact_autowrap.txt"
OUTPUT_DOCX = ROOT / "LastContact_new2.docx"
COLOR_PALETTE = [
    RGBColor(126, 72, 0),
    RGBColor(147, 69, 83),
    RGBColor(103, 84, 0),
    RGBColor(0, 102, 0),
    RGBColor(43, 43, 0),
    RGBColor(0, 93, 93),
    RGBColor(92, 53, 74),
]
HEADING_PATTERN = re.compile(r"^Часть\s+\S+$", re.I)
EPIGRAPH = "Шум — это не мусор. Это то, что пока не нашло своего имени."

TABLE_SEPARATOR_PATTERN = re.compile(r"^-+|:-+:|:-+|-[^|]*$")
IMAGE_LINE = re.compile(r"^\s*<img\s+([^>]*?)\s*\/?>\s*$", re.I)
ATTR = re.compile(r"([^\s=\/<>{}]+)\s*=\s*(?:\"([^\"]*)\"|'([^']*)'|([^\s'\"=<>`]+))")


def parse_image_line(line: str):
    m = IMAGE_LINE.match(line)
    if not m:
        return None
    attrs = {}
    for name, dq, sq, bare in ATTR.findall(m.group(1)):
        attrs[name.lower()] = dq or sq or bare or ""
    return attrs.get("src", "").strip() or None


def parse_table_row(line: str):
    txt = line.strip()
    if not (txt.startswith("|") and txt.endswith("|")):
        return None
    cells = [c.strip() for c in txt[1:-1].split("|")]
    return cells if len(cells) >= 2 else None


def is_table_separator(line: str) -> bool:
    txt = line.strip()
    if not txt.startswith("|"):
        return False
    cells = [c.replace("—", "-").strip() for c in txt[1:-1].split("|")]
    if not cells:
        return False
    return all(TABLE_SEPARATOR_PATTERN.fullmatch(cell) is not None for cell in cells)


def parse_blocks(lines):
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip() == "---":
            blocks.append(("text", ""))
            i += 1
            continue

        src = parse_image_line(line)
        if src:
            blocks.append(("image", src))
            i += 1
            continue

        table_row = parse_table_row(line)
        if table_row is not None and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [table_row]
            i += 2
            while i < len(lines):
                candidate = parse_table_row(lines[i])
                if candidate is None or is_table_separator(lines[i]):
                    break
                rows.append(candidate)
                i += 1

            if len(rows) >= 2:
                blocks.append(("table", rows))
                continue

            blocks.append(("text", line))
            continue

        blocks.append(("text", line))
        i += 1

    return blocks or [("text", "")]


def clear_document(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_document_background(document: Document) -> None:
    """Удаляет специальные свойства фона, оставляя стандартный фон Word."""
    document_elm = document._element
    for child in list(document_elm):
        if child.tag == qn("w:background"):
            document_elm.remove(child)

    settings_elm = document.settings._element
    for child in list(settings_elm):
        if child.tag == qn("w:displayBackgroundShape"):
            settings_elm.remove(child)

    normal_style = document.styles["Normal"]._element
    p_pr = normal_style.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn("w:shd"):
            p_pr.remove(child)


def set_paragraph_background(paragraph):
    """Удаляет локальную заливку, оставляя стандартный фон абзаца."""
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn("w:shd"):
            p_pr.remove(child)


def is_heading(text: str) -> bool:
    stripped = text.strip()
    return stripped in {"МИХАИЛ КРАВЧЕНКО", "ПОСЛЕДНИЙ КОНТАКТ"} or bool(
        HEADING_PATTERN.fullmatch(stripped)
    )


def add_text_line(document: Document, text: str, color=None):
    p = document.add_paragraph()
    p.style = document.styles["Normal"]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    if is_heading(text):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif text.strip() == EPIGRAPH:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    m = re.match(r"^[ \t]+", text)
    prefix = ""
    if m:
        raw = m.group(0)
        text = text[len(raw):]
        prefix = raw.replace(" ", "\u00A0").replace("\t", "\u00A0\u00A0\u00A0\u00A0")

    run = p.add_run(prefix + text)
    if color is not None:
        run.font.color.rgb = color

    set_paragraph_background(p)


def add_table(document: Document, rows, color_index):
    table = document.add_table(rows=0, cols=max(len(r) for r in rows))
    table.style = "Normal Table"
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        row_color = COLOR_PALETTE[(color_index + row_idx) % len(COLOR_PALETTE)]
        for col_idx, value in enumerate(row):
            cell = cells[col_idx]
            p = cell.paragraphs[0]
            p.clear()
            set_paragraph_background(p)
            run = p.add_run(value)
            run.font.color.rgb = row_color
            if row_idx == 0:
                run.bold = True
    return color_index + len(rows)


def add_image(document: Document, source_name: str):
    image_path = ROOT / source_name
    if not image_path.exists():
        return

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.95))
    set_paragraph_background(p)


def build_docx():
    text = TXT_PATH.read_text(encoding="utf-8").replace("\r\n", "\n").split("\n")
    blocks = parse_blocks(text)

    template = Document(str(TEMPLATE_DOCX))
    clear_document(template)
    set_document_background(template)

    color_index = 0
    for kind, data in blocks:
        if kind == "text":
            color = None
            if data.strip():
                color = COLOR_PALETTE[color_index % len(COLOR_PALETTE)]
                color_index += 1
            add_text_line(template, data, color)
        elif kind == "image":
            add_image(template, data)
        else:
            color_index = add_table(template, data, color_index)

    template.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
